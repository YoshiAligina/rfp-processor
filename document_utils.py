"""
Document Processing Utilities for RFP Analysis System

This module provides comprehensive text extraction capabilities for multiple document
formats commonly used in RFP submissions. It handles PDF, DOCX, and Excel files with
robust error handling and fallback mechanisms. The module integrates with OCR
capabilities for scanned documents and provides a unified interface for extracting
readable text from various file types.

Key features:
- Multi-format document support (PDF, DOCX, Excel)
- OCR integration for scanned documents
- Structured text extraction from tables and forms
- MIME type detection for web serving
- Robust error handling and fallback mechanisms
"""

import os
from docx import Document
import pandas as pd
import openpyxl
from ocr_utils import is_scanned, extract_text as extract_pdf_text

def extract_text_from_docx(file_path):
    """
    Extracts comprehensive text content from Microsoft Word DOCX documents with background processing support.
    
    BACKGROUND PROCESSING FEATURES:
    - Progress output flushing for real-time status updates
    - Memory-efficient processing of large documents
    - Robust error handling to prevent crashes
    - Immediate feedback during long operations
    
    This function processes both paragraph text and table data from Word documents,
    providing complete text extraction for RFP analysis. It handles document
    structure by extracting paragraphs sequentially and processing tables with
    proper cell delimitation. Essential for analyzing Word-based RFP submissions
    which often contain structured information in both narrative and tabular formats.
    
    Args:
        file_path (str): Full path to the DOCX file to process
        
    Returns:
        str: Extracted text content with paragraphs and table data
    """
    import sys
    
    try:
        print(f"[BACKGROUND] Starting DOCX extraction: {file_path}")
        sys.stdout.flush()
        
        doc = Document(file_path)
        text_content = []
        
        # Extract text from paragraphs with progress indication
        print(f"[BACKGROUND] Extracting paragraphs from DOCX...")
        sys.stdout.flush()
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables with progress indication
        print(f"[BACKGROUND] Extracting tables from DOCX...")
        sys.stdout.flush()
        
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        result = "\n".join(text_content)
        print(f"[BACKGROUND] DOCX extraction completed: {len(result)} characters")
        sys.stdout.flush()
        return result
        
    except Exception as e:
        error_msg = f"Error extracting text from DOCX: {e}"
        print(f"[BACKGROUND ERROR] {error_msg}")
        sys.stdout.flush()
        return ""

def extract_text_from_excel(file_path):
    """
    Extracts structured text content from Excel spreadsheets with multi-sheet support and background processing.
    
    BACKGROUND PROCESSING FEATURES:
    - Real-time progress updates during sheet processing
    - Memory-efficient handling of large spreadsheets
    - Immediate output flushing for status visibility
    - Robust error handling with fallback methods
    
    This function processes Excel files (.xlsx and .xls) by reading all worksheets
    and converting tabular data into readable text format. It uses openpyxl as the
    primary extraction method with pandas as fallback for broader compatibility.
    Particularly useful for RFP submissions that include pricing sheets, vendor
    information, or technical specifications in spreadsheet format. Maintains
    sheet structure and cell relationships in the extracted text.
    
    Args:
        file_path (str): Full path to the Excel file to process
        
    Returns:
        str: Extracted text content from all worksheets with sheet labels
    """
    import sys
    
    try:
        print(f"[BACKGROUND] Starting Excel extraction: {file_path}")
        sys.stdout.flush()
        
        # Try reading with openpyxl first (for .xlsx)
        try:
            print(f"[BACKGROUND] Using openpyxl for Excel processing...")
            sys.stdout.flush()
            
            workbook = openpyxl.load_workbook(file_path, data_only=True)
            text_content = []
            
            print(f"[BACKGROUND] Processing {len(workbook.sheetnames)} Excel sheets...")
            sys.stdout.flush()
            
            for sheet_name in workbook.sheetnames:
                print(f"[BACKGROUND] Processing sheet: {sheet_name}")
                sys.stdout.flush()
                
                sheet = workbook[sheet_name]
                text_content.append(f"Sheet: {sheet_name}")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = []
                    for cell in row:
                        if cell is not None and str(cell).strip():
                            row_text.append(str(cell).strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            result = "\n".join(text_content)
            print(f"[BACKGROUND] Excel extraction completed: {len(result)} characters")
            sys.stdout.flush()
            return result
            
        except Exception as openpyxl_error:
            # Fallback to pandas for both .xlsx and .xls
            print(f"[BACKGROUND] Openpyxl failed, trying pandas fallback...")
            sys.stdout.flush()
            
            excel_file = pd.ExcelFile(file_path)
            text_content = []
            
            print(f"[BACKGROUND] Pandas processing {len(excel_file.sheet_names)} sheets...")
            sys.stdout.flush()
            
            for sheet_name in excel_file.sheet_names:
                print(f"[BACKGROUND] Processing sheet with pandas: {sheet_name}")
                sys.stdout.flush()
                
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text_content.append(f"Sheet: {sheet_name}")
                
                # Convert DataFrame to text
                for _, row in df.iterrows():
                    row_text = []
                    for value in row:
                        if pd.notna(value) and str(value).strip():
                            row_text.append(str(value).strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            result = "\n".join(text_content)
            print(f"[BACKGROUND] Pandas Excel extraction completed: {len(result)} characters")
            sys.stdout.flush()
            return result
            
    except Exception as e:
        error_msg = f"Error extracting text from Excel: {e}"
        print(f"[BACKGROUND ERROR] {error_msg}")
        sys.stdout.flush()
        return ""

def get_file_type(filename):
    """
    Determines document type from file extension for processing workflow routing.
    This function analyzes file extensions to categorize documents into supported
    types (PDF, DOCX, Excel) enabling the system to route files to appropriate
    text extraction functions. Essential for the document processing pipeline
    as it determines which extraction method and OCR capabilities to apply.
    Supports both modern and legacy Excel formats for broader compatibility.
    
    Args:
        filename (str): Name of the file including extension
        
    Returns:
        str: File type identifier ('pdf', 'docx', 'excel', or 'unknown')
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext == '.pdf':
        return 'pdf'
    elif ext == '.docx':
        return 'docx'
    elif ext in ['.xlsx', '.xls']:
        return 'excel'
    else:
        return 'unknown'

def extract_text_from_file(file_path, use_ocr=False):
    """
    Universal text extraction function that handles multiple document formats intelligently.
    This is the main entry point for document text extraction, providing a unified
    interface that automatically detects file type and applies appropriate extraction
    methods. For PDFs, it intelligently determines whether OCR is needed for scanned
    documents. Integrates all specialized extraction functions into a single,
    easy-to-use interface that handles the complete text extraction workflow
    with optimal method selection based on document characteristics.
    
    Args:
        file_path (str): Full path to the document file
        use_ocr (bool): Force OCR usage (mainly for PDF processing)
        
    Returns:
        str: Extracted text content ready for ML analysis
    """
    file_type = get_file_type(file_path)
    
    if file_type == 'pdf':
        if is_scanned(file_path):
            return extract_pdf_text(file_path, use_easyocr=True)
        else:
            return extract_pdf_text(file_path)
    elif file_type == 'docx':
        return extract_text_from_docx(file_path)
    elif file_type == 'excel':
        return extract_text_from_excel(file_path)
    else:
        return ""

def get_appropriate_mime_type(filename):
    """
    Determines the correct MIME type for document files to enable proper web serving.
    This function maps file extensions to their corresponding MIME types, essential
    for web applications that need to serve documents with correct content headers.
    Supports the major document formats used in RFP processing and provides a
    generic fallback for unknown formats. Used primarily for file downloads
    and web browser compatibility in the Flask web interface.
    
    Args:
        filename (str): Name of the file including extension
        
    Returns:
        str: Appropriate MIME type string for the file format
    """
    ext = os.path.splitext(filename)[1].lower()
    mime_types = {
        '.pdf': 'application/pdf',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        '.xls': 'application/vnd.ms-excel'
    }
    return mime_types.get(ext, 'application/octet-stream')
